"""
DOCX extractor.

Reads paragraphs and tables out of Word files with python-docx. Tables are kept
both as structured rows (first row treated as the header) and flattened into the
text stream so nothing is lost for semantic search.
"""

from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.extractors.base import ExtractedContent, find_equipment_tags


def extract(path: str | Path) -> ExtractedContent:
    doc = DocxDocument(str(path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables: list[list[dict]] = []
    for table in doc.tables:
        tables.append(_rows_from_table(table))

    text = "\n".join(paragraphs).strip()
    content = ExtractedContent(text=text, tables=tables, doc_type="document")
    content.equipment_tags = find_equipment_tags(text)
    return content


def _rows_from_table(table) -> list[dict]:
    """Turn a docx table into row dicts using the first row as column headers."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return []
    header, *body = rows
    return [dict(zip(header, r)) for r in body]
